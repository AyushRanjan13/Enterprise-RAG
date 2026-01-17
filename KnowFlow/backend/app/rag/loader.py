"""
Multi-format document loader for PDF, DOCX, and TXT files.
Extracts text and maintains metadata.
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
import logging
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class MultiFormatLoader:
    """Loads documents from multiple file formats with metadata."""
    
    SUPPORTED_FORMATS = {"pdf", "docx", "txt"}
    
    def __init__(self):
        """Initialize the loader with optional dependencies check."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and warn about optional dependencies."""
        try:
            import pypdf
        except ImportError:
            logger.warning("pypdf not installed. PDF support disabled.")
        
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed. DOCX support disabled.")
    
    def load_file(
        self,
        file_path: str,
        metadata: Optional[Dict] = None
    ) -> List[Document]:
        """
        Load a single file and return documents with metadata.
        
        Args:
            file_path: Path to the file
            metadata: Additional metadata to attach
        
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower().lstrip(".")
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Loading file: {file_path}")
        
        if file_ext == "pdf":
            return self._load_pdf(file_path, metadata)
        elif file_ext == "docx":
            return self._load_docx(file_path, metadata)
        elif file_ext == "txt":
            return self._load_txt(file_path, metadata)
    
    def _load_pdf(
        self,
        file_path: Path,
        metadata: Optional[Dict] = None
    ) -> List[Document]:
        """Load PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf required for PDF support. Install with: pip install pypdf")
        
        documents = []
        reader = PdfReader(str(file_path))
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            
            doc_metadata = {
                "source": str(file_path.name),
                "page_number": page_num,
                **(metadata or {})
            }
            
            doc = Document(
                page_content=text,
                metadata=doc_metadata
            )
            documents.append(doc)
        
        logger.info(f"Loaded {len(documents)} pages from PDF: {file_path.name}")
        return documents
    
    def _load_docx(
        self,
        file_path: Path,
        metadata: Optional[Dict] = None
    ) -> List[Document]:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required for DOCX support. Install with: pip install python-docx")
        
        docx = DocxDocument(str(file_path))
        text = "\n".join([para.text for para in docx.paragraphs])
        
        doc_metadata = {
            "source": str(file_path.name),
            **(metadata or {})
        }
        
        doc = Document(
            page_content=text,
            metadata=doc_metadata
        )
        
        logger.info(f"Loaded DOCX: {file_path.name}")
        return [doc]
    
    def _load_txt(
        self,
        file_path: Path,
        metadata: Optional[Dict] = None
    ) -> List[Document]:
        """Load TXT file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        doc_metadata = {
            "source": str(file_path.name),
            **(metadata or {})
        }
        
        doc = Document(
            page_content=text,
            metadata=doc_metadata
        )
        
        logger.info(f"Loaded TXT: {file_path.name}")
        return [doc]
    
    def load_directory(
        self,
        dir_path: str,
        metadata: Optional[Dict] = None,
        recursive: bool = True
    ) -> List[Document]:
        """
        Load all supported files from a directory.
        
        Args:
            dir_path: Directory path
            metadata: Additional metadata
            recursive: Whether to search subdirectories
        
        Returns:
            List of all loaded documents
        """
        dir_path = Path(dir_path)
        documents = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    docs = self.load_file(file_path, metadata)
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {str(e)}")
        
        logger.info(f"Loaded {len(documents)} documents from {dir_path}")
        return documents
